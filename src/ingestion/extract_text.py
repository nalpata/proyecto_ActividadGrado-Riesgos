from pathlib import Path
import re
import fitz
from docx import Document

RAW_DIR = Path("data/raw")
OUT_DIR = Path("data/processed/text")
OUT_DIR.mkdir(parents=True, exist_ok=True)

def safe_name(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", Path(name).stem) + ".txt"

def extract_pdf(path: Path) -> str:
    doc = fitz.open(path)
    return "\n".join(page.get_text("text") for page in doc)

def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)

def main():
    files = list(RAW_DIR.rglob("*.pdf")) + list(RAW_DIR.rglob("*.docx"))
    for file in files:
        text = extract_pdf(file) if file.suffix.lower() == ".pdf" else extract_docx(file)
        out = OUT_DIR / safe_name(file.name)
        out.write_text(text, encoding="utf-8")
        print(f"OK: {file.name} -> {out}")

if __name__ == "__main__":
    main()
